#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工具：从 .docx/.xlsx 里抽取纯文本内容（移植自 PTA 的 tools/office_text.py，逻辑不变）。

批量概念笔记提炼的候选文档包含 .docx，需要同样的抽取能力——这两种格式本质是
zip+XML，不能直接当文本 decode。每个 Agent 的 tools/ 自包含，不跨 Agent
import，同 llm_client.py 的移植先例。
"""

from pathlib import Path


def extract_docx_text(path: Path) -> str:
    try:
        import docx
    except ImportError:
        return "[未安装 python-docx，无法解析该 .docx 文件内容，仅能感知到文件本身发生了变化]"
    try:
        document = docx.Document(str(path))
    except Exception as e:
        return f"[无法解析该 .docx 文件: {e}]"

    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx_text(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return "[未安装 openpyxl，无法解析该 .xlsx 文件内容，仅能感知到文件本身发生了变化]"
    try:
        workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    except Exception as e:
        return f"[无法解析该 .xlsx 文件: {e}]"

    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                cells = [str(cell) if cell is not None else "" for cell in row]
                parts.append(" | ".join(cells))
    return "\n".join(parts)


OFFICE_EXTRACTORS = {
    ".docx": extract_docx_text,
    ".xlsx": extract_xlsx_text,
}


def extract_office_text(path: Path) -> str:
    extractor = OFFICE_EXTRACTORS.get(Path(path).suffix.lower())
    if extractor is None:
        return ""
    return extractor(Path(path))
