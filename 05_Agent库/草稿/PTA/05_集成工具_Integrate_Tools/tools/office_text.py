#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工具：从 .docx/.xlsx 里抽取纯文本内容，供 daily_sensing 的逐行 diff 用。

这两种格式本质是 zip 压缩包（里面套 XML），不能像 .md/.py 那样直接当文本
读——之前 DEFAULT_SCAN_EXTENSIONS 没覆盖它们就是因为这个限制：直接
decode 会是乱码或直接报错，喂给 LLM 也做不出有意义的摘要。这里只抽取
"读起来像纯文本"的内容（段落文字/表格文字/单元格值），不还原格式、图片、
批注——够 LLM 做语义摘要就行，不追求完整还原整份文档。

依赖 python-docx / openpyxl，是本项目目前唯一的两个非标准库依赖（其余
全是 stdlib）。如果运行环境没装，优雅降级：返回一句提示文字，不让整个
--daily-scan 因为这两个库缺失而崩掉——同一条"缺配置/缺依赖不该让主流程
崩溃"的原则，跟 wecom_notify.load_wecom_config 遇文件不存在返回 None
是一样的处理方式。
"""

from pathlib import Path


def extract_docx_text(path: Path) -> str:
    try:
        import docx
    except ImportError:
        return "[未安装 python-docx，无法解析该 .docx 文件内容，仅能感知到文件本身发生了变化]"
    try:
        document = docx.Document(str(path))
    except Exception as e:
        return f"[无法解析该 .docx 文件: {e}]"

    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx_text(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return "[未安装 openpyxl，无法解析该 .xlsx 文件内容，仅能感知到文件本身发生了变化]"
    try:
        # data_only=True：读公式的最后计算结果而不是公式本身（对"内容变了没有"
        # 这个判断更有意义）；read_only=True：大表格也不会整个加载进内存。
        workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    except Exception as e:
        return f"[无法解析该 .xlsx 文件: {e}]"

    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                cells = [str(cell) if cell is not None else "" for cell in row]
                parts.append(" | ".join(cells))
    return "\n".join(parts)


#: 后缀 → 抽取函数。daily_sensing.py 用这个字典的 key 判断"这个文件需要走
#: 抽取而不是直接当文本读"，不需要在两处分别维护一份后缀列表。
OFFICE_EXTRACTORS = {
    ".docx": extract_docx_text,
    ".xlsx": extract_xlsx_text,
}


def extract_office_text(path: Path) -> str:
    extractor = OFFICE_EXTRACTORS.get(Path(path).suffix.lower())
    if extractor is None:
        return ""
    return extractor(Path(path))
